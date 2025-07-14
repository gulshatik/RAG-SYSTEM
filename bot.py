import os
import re
import json
from docx import Document
import chromadb
from chromadb.utils import embedding_functions
import hashlib
import win32com.client
from openai import OpenAI
import time
import logging
import traceback
import sys
from docx.shared import Pt
import PyPDF2
import textwrap
from collections import deque
from vector_rag_db import VectorRAGDatabase
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
    ConversationHandler
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)
logger.info("Логгирование настроено с кодировкой UTF-8")
POSITION, DEPARTMENT = range(2)
DOCUMENTS_DIR = r"ПОЛНЫЙ ПУТЬ К ПАПКЕ С ГОТОВЫМИ ДОЛЖНОСТНЫМИ ИНСТРУКЦИЯМИ"
VECTOR_DB_PATH = r"ПОЛНЫЙ ПУТЬ К ПАПКЕ ГДЕ БУДЕТ ХРАНИТЬСЯ ВЕКТОРНАЯ БАЗА ДАННЫХ"
OUTPUT_DIR = r"ПОЛНЫЙ ПУТЬ К ПАПКЕ ГДЕ БУДУТ ХРАНИТЬСЯ РЕЗУЛЬТАТЫ"
TEMPLATE_PATH = r"ПОЛНЫЙ ПУТЬ К ДОКУМЕНТУ Шаблон.docx"
TOKEN = "ТОКЕН ТЕЛЕГРАМ БОТА"

LAST_API_REQUEST_TIME = 0
vector_db = None
deepseek_client = None

def init_system():
    """Инициализация системных компонентов"""
    global vector_db, deepseek_client
    
    deepseek_client = OpenAI(
        base_url="ССЫЛКА SCALEWAY",
        api_key="API КЛЮЧ"
    )
    
    vector_db = VectorRAGDatabase(DOCUMENTS_DIR, VECTOR_DB_PATH)
    logger.info("Системные компоненты инициализированы")

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начало диалога"""
    await update.message.reply_text(
        "Привет! Я бот для генерации должностных инструкций.\n"
        "Введите название должности:",
        reply_markup=ReplyKeyboardRemove()
    )
    return POSITION

async def get_position(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получение названия должности"""
    context.user_data['position'] = update.message.text
    await update.message.reply_text("Теперь введите название подразделения:")
    return DEPARTMENT

async def get_department(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получение названия подразделения и запуск генерации"""
    context.user_data['department'] = update.message.text
    position = context.user_data['position']
    department = context.user_data['department']
    
    await update.message.reply_text(
        f"Начинаю генерацию инструкции для:\n"
        f"Должность: {position}\n"
        f"Подразделение: {department}\n"
        "Это займет несколько минут..."
    )
    try:
        output_path = generate_job_description(position, department)
        with open(output_path, 'rb') as doc_file:
            await update.message.reply_document(
                document=doc_file,
                caption=f"Должностная инструкция для {position}"
            )
        os.remove(output_path)  
    except Exception as e:
        logger.error(f"Ошибка генерации: {str(e)}")
        await update.message.reply_text("Произошла ошибка при генерации документа 😢")
    
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Отмена диалога"""
    await update.message.reply_text(
        "Генерация отменена",
        reply_markup=ReplyKeyboardRemove()
    )
    return ConversationHandler.END

def generate_job_description(position: str, department: str) -> str:
    """Генерация документа (адаптированная версия вашей main)"""
    global LAST_API_REQUEST_TIME
    
    template_doc = Document(TEMPLATE_PATH)
    if not template_doc:
        raise Exception("Не удалось загрузить шаблон")
    
    processed_doc = process_template(template_doc, position, department)
    if not processed_doc:
        raise Exception("Ошибка обработки шаблона")
    
    output_filename = f"ДИ_{position}_{time.strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    if not save_document(processed_doc, output_path):
        raise Exception("Ошибка сохранения документа")
    
    return output_path

def to_accusative_via_llm(phrase: str) -> str:
    """Преобразует фразу в родительный падеж с помощью LLM"""
    global LAST_API_REQUEST_TIME
    try:
        logger.info(f"Запрос к LLM для преобразования в родительный падеж: {phrase}")


        current_time = time.time()
        if LAST_API_REQUEST_TIME > 0:
            elapsed = current_time - LAST_API_REQUEST_TIME
            if elapsed < 60:  
                sleep_time = 60 - elapsed
                logger.info(f"Ожидание {sleep_time:.2f} сек перед запросом...")
                time.sleep(sleep_time)
        
        LAST_API_REQUEST_TIME = time.time()

        prompt = (
            "<think>\n"
            f"Преобразую название должности '{phrase}' в родительный падеж.\n"
            "Анализирую структуру должности и применяю правила русского языка.\n"
            "Определяю правильные окончания для каждого слова в должности.\n"
            "</think>\n"
            f"родительный падеж: "
        )
        
        response = deepseek_client.chat.completions.create(
            model="deepseek-r1-distill-llama-70b",
            messages=[
                {
                    "role": "system", 
                    "content": (
                        "Ты эксперт по русскому языку. Всегда используй формат с тегами <think> для размышлений. "
                        "и после тега </think> предоставляй только окончательный ответ - преобразованное название должности. "
                        "Используй нижний регистр и не добавляй никаких дополнительных слов или комментариев.ответ должен быть в формате </think> 'ответ' <end>"
                    )
                },
                {"role": "user", "content": prompt}
            ],
            max_tokens=50,
            temperature=0.0,
            stop=["<end>"] 
        )
        
        result = response.choices[0].message.content.strip()
        
        if "</think>" in result:
            generated_text = result.split("</think>")[-1].strip()
            
            prefixes = ["родительный падеж:", "Ответ:", "Результат:", "'ответ'", "•", "-", "'"]
            for prefix in prefixes:
                if generated_text.startswith(prefix):
                    generated_text = generated_text[len(prefix):].strip()
            
            if generated_text.startswith('"') and generated_text.endswith('"'):
                generated_text = generated_text[1:-1].strip()
            elif generated_text.startswith("'") and generated_text.endswith("'"):
                generated_text = generated_text[1:-1].strip()
                
            logger.info(f"Преобразовано в родительный падеж: {phrase} -> {generated_text}")
            return generated_text
        else:
            clean_result = re.sub(r'</?[a-z]+>', '', result, flags=re.IGNORECASE)
            logger.info(f"Преобразовано без тега: {phrase} -> {clean_result}")
            return clean_result
    
    except Exception as e:
        logger.error(f"Ошибка преобразования в родительный падеж: {str(e)}")
        return phrase 

PLACEHOLDER_CONFIG = {
    "ОБЩИЕ ПОЛОЖЕНИЯ 1.2": {
        "prompt": "Составь пункт 1.2 'Требования к образованию и стажу работы' раздела 'Общие положения' для должности {position} в подразделении {department}. Используй юридические формулировки, используй ТОЛЬКО русский язык. 1-2 предложения",
        "context_query": "требования к образованию и стажу работы для должности {position}"
    },
    "ОБЩИЕ ПОЛОЖЕНИЯ 1.3": {
        "prompt": "Составь пункт 1.3 раздела 'Общие положения' для {position} ({department}). Ответ должен содержать только текст 'Лицо принимается на должность {position} и освобождается от нее приказом ректора КФУ (проректора, иного уполномоченного ректором лица) в установленном действующим законодательством Российской Федерации порядке. {position} подчиняется непосредственно (тут напиши 'главе {department}').' больше никакого текста.",
        "context_query": "порядок приема и освобождения от должности {position}"
    },
    "ОБЩИЕ ПОЛОЖЕНИЯ 1.4": {
        "prompt": "Перечисли что должен знать {position} ({department}). Формат: - [текст]. Ответ обязательно должен содержать текст '- Конституцию Российской Федерации, иные законодательные и нормативные правовые акты, методические материалы, определяющие направления развития высшего образования в Российской Федерации, а также нормативные правовые акты, методические материалы по тематике работы, основы трудового законодательства и организации труда в РФ;- Устав КФУ, Правила внутреннего распорядка КФУ, Положение о пропускном и внутриобъектовом режиме КФУ, Коллективный договор КФУ, Положение о (полное наименование структурного подразделения), Положение о (полное наименование лаборатории), Кодекс этики и служебного поведения работников КФУ и иные локальные нормативные акты КФУ;- структуру и специфику деятельности КФУ;'Далее добавь 3-12 пунктов (в зависимости от должности) какими знаниями должен обладать сотрудник. Используй юридические формулировки, используй ТОЛЬКО русский язык, никакого дополнительного текста.",
        "context_query": "знания и компетенции для должности {position}"
    },
    "ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ": {
        "prompt": "Какие должностные обязанности исполняет {position} ({department}). Формат: - [текст]. Любое количество пунктов, в каждом пункте максимум 2 предложения. Используй юридические формулировки, используй ТОЛЬКО русский язык, никакого дополнительного текста.",
        "context_query": "должностные обязанности {position}"
    }
}

def generate_placeholder_content(placeholder: str, context: dict) -> str:
    """Генерация содержания с использованием RAG"""
    global LAST_API_REQUEST_TIME
    try:
        logger.info(f"Генерация для плейсхолдера: [{placeholder}]")
        
        if placeholder not in PLACEHOLDER_CONFIG:
            logger.error(f"Неизвестный плейсхолдер: {placeholder}")
            return f"[{placeholder}]"
        
        config = PLACEHOLDER_CONFIG[placeholder]
        prompt_template = config["prompt"]
        context_query = config.get("context_query", "").format(**context)

        base_prompt = prompt_template.format(**context)

        logger.info(f"Промпт для генерации: {base_prompt[:200]}...")

        if context_query:
            relevant_chunks = vector_db.search_relevant_chunks(context_query, n_results=3)
            rag_context = "\n\n".join([f"Источник: {chunk['source']}\nКонтент: {chunk['content']}" 
                                      for chunk in relevant_chunks])
            
            full_prompt = (
                "Используй следующие фрагменты из должностных инструкций "
                "и нормативных документов как контекст для ответа:\n\n"
                f"{rag_context}\n\n"
                "Задание:\n"
                f"{base_prompt}"
            )
        else:
            full_prompt = base_prompt
        
        logger.info(f"Промпт для генерации: {full_prompt[:500]}...")

        current_time = time.time()
        if LAST_API_REQUEST_TIME > 0:
            elapsed = current_time - LAST_API_REQUEST_TIME
            if elapsed < 60: 
                sleep_time = 60
                logger.info(f"Ожидание {sleep_time:.2f} сек перед запросом...")
                time.sleep(sleep_time)
        
        LAST_API_REQUEST_TIME = time.time()

        response = deepseek_client.chat.completions.create(
            model="deepseek-r1-distill-llama-70b",
            messages=[
                {
                    "role": "system", 
                    "content": (
                        "Ты HR-специалист, составляющий должностные инструкции. "
                        "Используй юридические формулировки и только русский язык."
                    )
                },
                {"role": "user", "content": full_prompt}
            ],
            max_tokens=1500,
            temperature=0.3
        )
        
        result = response.choices[0].message.content.strip()
        if "</think>" in result:

            generated_text = result.split("</think>")[-1].strip()

            generated_text = re.sub(r'</?[a-z]+>', '', generated_text, flags=re.IGNORECASE)
            logger.info(f"Отфильтрованный контент: {generated_text[:100]}...")
            return generated_text
        else:
            logger.info(f"Контент без фильтрации: {result[:100]}...")
            return result
    
    except Exception as e:
        logger.error(f"Ошибка генерации: {str(e)}")
        logger.error(traceback.format_exc())
        return f"[{placeholder}]"          


def read_docx(file_path: str) -> Document:
    """Чтение .docx файла"""
    try:
        logger.info(f"Чтение .docx файла: {os.path.basename(file_path)}")
        return Document(file_path)
    except Exception as e:
        logger.error(f"Ошибка чтения DOCX: {str(e)}")
        return None

def process_template(template_doc: Document, position: str, department: str) -> Document:
    """Обработка шаблона с заменой плейсхолдеров"""
    try:
        logger.info("Начало обработки шаблона...")
        

        context = {
            "position": position,
            "department": department
        }
        
        placeholder_pattern = re.compile(r"\[([^\]]+)\]")
        
        for paragraph in template_doc.paragraphs:
            if placeholder_pattern.search(paragraph.text):
                def replace_placeholder(match):
                    placeholder_name = match.group(1).strip()
                    
                    if "наименование должности" in placeholder_name.lower():
                        is_uppercase = placeholder_name[0].isupper()
                        return position if is_uppercase else to_accusative_via_llm(position)
                            
                    elif "наименование кафедры" in placeholder_name.lower():
                        return department
                    elif "наименование структурного подразделения" in placeholder_name.lower():
                        return department
                    
                    return generate_placeholder_content(placeholder_name, context)
                
                paragraph.text = placeholder_pattern.sub(replace_placeholder, paragraph.text)
        
        logger.info("Шаблон успешно обработан")
        return template_doc
    
    except Exception as e:
        logger.error(f"Ошибка обработки шаблона: {str(e)}")
        logger.error(traceback.format_exc())
        return None

def save_document(doc: Document, output_path: str) -> bool:
    """Сохранение документа"""
    try:
        logger.info(f"Сохранение документа: {output_path}")
        doc.save(output_path)
        return True
    except Exception as e:
        logger.error(f"Ошибка сохранения: {str(e)}")
        return False


def main() -> None:
    """Запуск бота"""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    init_system()
    
    application = Application.builder().token(TOKEN).build()
    
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            POSITION: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_position)],
            DEPARTMENT: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_department)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )
    
    application.add_handler(conv_handler)
    
    logger.info("Бот запущен")
    application.run_polling()

if __name__ == "__main__":
    main()
